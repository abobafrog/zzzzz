from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


TABLE_MIN_ROWS = 2


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).replace("\xa0", " ").split()).strip()



def table_to_rows(table) -> tuple[list[str], list[dict[str, str]]] | None:
    raw_rows: list[list[str]] = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        if any(cells):
            raw_rows.append(cells)

    if len(raw_rows) < TABLE_MIN_ROWS:
        return None

    header = raw_rows[0]
    if not any(header):
        return None

    columns = [col if col else f"column{idx + 1}" for idx, col in enumerate(header)]
    rows: list[dict[str, str]] = []

    for raw in raw_rows[1:]:
        padded = raw + [""] * max(0, len(columns) - len(raw))
        row_obj = {columns[idx]: padded[idx] for idx in range(len(columns))}
        if any(v != "" for v in row_obj.values()):
            rows.append(row_obj)

    return columns, rows



def parse_docx(path: str | Path) -> dict[str, Any]:
    doc = Document(str(path))

    warnings: list[str] = []
    paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    tables_data: list[tuple[list[str], list[dict[str, str]]]] = []

    for table in doc.tables:
        parsed = table_to_rows(table)
        if parsed:
            tables_data.append(parsed)

    if tables_data:
        columns, rows = tables_data[0]
        if len(tables_data) > 1:
            warnings.append(
                f"Found {len(tables_data)} tables in DOCX, returned only the first table."
            )
        return {
            "file_name": Path(path).name,
            "file_type": "docx",
            "content_type": "table",
            "columns": columns,
            "rows": rows,
            "text": "\n".join(paragraphs),
            "blocks": [{"type": "paragraph", "text": text} for text in paragraphs],
            "warnings": warnings,
        }

    warnings.append("No tables found in DOCX. Returned text blocks only.")
    return {
        "file_name": Path(path).name,
        "file_type": "docx",
        "content_type": "text",
        "columns": [],
        "rows": [],
        "text": "\n".join(paragraphs),
        "blocks": [{"type": "paragraph", "text": text} for text in paragraphs],
        "warnings": warnings,
    }
